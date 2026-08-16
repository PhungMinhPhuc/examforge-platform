"""Điều phối cấp đề thi cho bộ ghi Word mới (`doctree/write/docx.py`).

Tách khỏi `doctree/write/docx.py` vì module đó chỉ lo phần "cây -> docx" của
MỘT câu; phần "gom câu/đánh số theo phần/câu dẫn/đề thi hoàn chỉnh" là một
tầng khác, giữ ở đây khớp cách `exporters/pdf_html/renderer.py::render_exam_html`
đã tách (cùng logic đếm/đánh số/xử lý `st`, chỉ khác định dạng đích).
"""
import io
import gc
import shutil

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

from .common import fmt_points, get_scoring_config
from doctree.adapt import question_to_rec, figures_by_id
from doctree.write.docx import setup_document, add_question, write_field
from doctree.read.tex import read_field


SECTION_LABELS = {
    "mc": "PHẦN I. Câu trắc nghiệm nhiều phương án lựa chọn. Thí sinh trả lời từ câu 1 đến câu {n}. Mỗi câu hỏi thí sinh chỉ chọn một phương án.",
    "tf": "PHẦN II. Câu trắc nghiệm đúng sai. Thí sinh trả lời từ câu 1 đến câu {n}. Trong mỗi ý a), b), c), d) ở mỗi câu hỏi, thí sinh chọn đúng hoặc sai.",
    "sa": "PHẦN III. Câu trắc nghiệm trả lời ngắn. Thí sinh trả lời từ câu 1 đến câu {n}.",
    "oe": "PHẦN IV. Câu tự luận. Thí sinh trả lời từ câu 1 đến câu {n}.",
}


def add_exam_body(doc, questions, show_answers=True, include_solution=True, math_batch=None):
    """Dựng toàn bộ thân đề (không header/footer/đáp án — Phase 3/4) vào
    `doc` đã sẵn sàng (qua `setup_document`). `questions` là danh sách phẳng
    lấy từ CSDL, câu con của `st` có `parent_id` trỏ về câu cha — ĐÚNG khớp
    cách `render_exam_html` (`pdf_html/renderer.py:251-300`) duyệt, không
    qua `group_units()` vì bản thân việc quét `parent_id` trực tiếp đã đủ
    đơn giản và đã có 2 nơi (bộ HTML, bộ LaTeX cũ) dùng chung đúng cách này.
    `math_batch` — xem `doctree/write/docx.py::add_inline`.
    """
    totals = {
        qt: sum(1 for q in questions if q.get("question_type") == qt)
        for qt in SECTION_LABELS
    }
    printed = {qt: False for qt in SECTION_LABELS}
    q_counter = 1
    processed_ids = set()

    for q in questions:
        q_id = q.get("id")
        if q_id in processed_ids:
            continue

        q_type = q.get("question_type")
        eff_type = q_type
        children = []
        if q_type == "st":
            children = [c for c in questions if c.get("parent_id") == q_id]
            if children:
                eff_type = children[0].get("question_type", "mc")

        if eff_type in printed and not printed[eff_type]:
            label = SECTION_LABELS[eff_type].format(n=totals[eff_type])
            # Chỉ tên phần được in đậm; câu hướng dẫn bắt đầu bằng "Thí sinh"
            # giữ kiểu thường, đồng bộ với preview/PDF.
            heading, separator, instruction = label.partition(" Thí sinh")
            p = doc.add_paragraph()
            p.add_run(heading).bold = True
            if separator:
                p.add_run(f" Thí sinh{instruction}")
            printed[eff_type] = True
            q_counter = 1

        if q_type == "st":
            figures = figures_by_id(q.get("images"))
            rec = question_to_rec(q)
            if children:
                start_c, end_c = q_counter, q_counter + len(children) - 1
                lead_text = f"Dựa vào thông tin dưới đây để trả lời các câu từ {start_c} đến {end_c}."
            else:
                lead_text = "Dựa vào thông tin dưới đây:"
            doc.add_paragraph().add_run(lead_text).italic = True
            write_field(doc, rec["content_doc"], figures, math_batch=math_batch)

            for child in children:
                child_figures = figures_by_id(child.get("images"))
                add_question(
                    doc, question_to_rec(child), q_counter, child_figures,
                    show_answers=show_answers, include_solution=include_solution,
                    math_batch=math_batch,
                )
                processed_ids.add(child.get("id"))
                q_counter += 1
            processed_ids.add(q_id)

        elif not q.get("parent_id"):
            figures = figures_by_id(q.get("images"))
            add_question(
                doc, question_to_rec(q), q_counter, figures,
                show_answers=show_answers, include_solution=include_solution,
                math_batch=math_batch,
            )
            processed_ids.add(q_id)
            q_counter += 1

    # Khớp nguyên xi `word_exporter.py:285`: `"-\/" * 21 + " HẾT " + "-\/" * 21`
    # — `\/` trong LaTeX là hiệu chỉnh nghiêng (kerning), không ra ký tự nhìn
    # thấy được, nên hiển thị thật ra đúng 21 dấu gạch mỗi bên.
    end_p = doc.add_paragraph()
    end_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    end_p.add_run("-" * 21 + " HẾT " + "-" * 21).bold = True


def _create_field(field_text, italic=False):
    """Field code Word thô (`NUMPAGES`, `PAGE`...) — port nguyên xi
    `word_exporter.py:433-471` (thuần OOXML, không phụ thuộc pandoc)."""
    def rpr():
        el = OxmlElement("w:rPr")
        if italic:
            el.append(OxmlElement("w:i"))
        return el

    fld1 = OxmlElement("w:fldChar")
    fld1.set(qn("w:fldCharType"), "begin")
    r1 = OxmlElement("w:r")
    r1.append(rpr())
    r1.append(fld1)

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = f" {field_text} "
    r2 = OxmlElement("w:r")
    r2.append(rpr())
    r2.append(instr)

    fld2 = OxmlElement("w:fldChar")
    fld2.set(qn("w:fldCharType"), "separate")
    r3 = OxmlElement("w:r")
    r3.append(rpr())
    r3.append(fld2)

    t = OxmlElement("w:t")
    t.text = "..."
    r4 = OxmlElement("w:r")
    r4.append(rpr())
    r4.append(t)

    fld3 = OxmlElement("w:fldChar")
    fld3.set(qn("w:fldCharType"), "end")
    r5 = OxmlElement("w:r")
    r5.append(rpr())
    r5.append(fld3)

    return [r1, r2, r3, r4, r5]


def _remove_borders(table):
    tbl_pr = table._tbl.tblPr
    for old in tbl_pr.findall(qn("w:tblBorders")):
        tbl_pr.remove(old)
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:val"), "none")
        borders.append(e)
    tbl_pr.append(borders)


def _insert_header_rule_after(cell, target, *, cell_width_cm, width_ratio):
    """Chèn đường gạch căn giữa trong flow Word, không dùng shape/anchor.

    Paragraph chứa đường kẻ cao đúng 1pt và được thụt đều hai phía để giữ
    đúng tỷ lệ với ô header. Vì nó là phần tử bình thường trong ô bảng nên
    luôn đi cùng header, không thể trôi hoặc nhảy sang vị trí khác.
    """
    rule = cell.add_paragraph()
    target._p.addnext(rule._p)
    side_indent = cell_width_cm * (1.0 - width_ratio) / 2.0
    fmt = rule.paragraph_format
    fmt.left_indent = Cm(side_indent)
    fmt.right_indent = Cm(side_indent)
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(0)
    fmt.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    fmt.line_spacing = Pt(1)

    p_pr = rule._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "0")
    bottom.set(qn("w:color"), "000000")
    borders.append(bottom)
    p_pr.append(borders)
    return rule


def add_page_break(doc):
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def add_exam_header(doc, code, exam_title, department, exam_type, subject, duration, is_answer=False):
    """2 bảng không viền + hộp "Mã đề" — port `build_header()`
    (`word_exporter.py:498-594`). Bản cũ phải chèn TRƯỚC một đoạn neo sẵn có
    (`anchor_p._p.addprevious`) vì nó hậu xử lý một file đã dựng xong qua
    pandoc; ở đây dựng tuần tự nên gọi hàm này ĐÚNG lúc muốn header xuất
    hiện (`doc.add_table()` tự nối vào cuối tài liệu) là đủ, không cần neo.
    Gọi được nhiều lần vào cùng 1 doc (đề sạch + trang đáp án).
    """
    t1 = doc.add_table(rows=1, cols=2)
    _remove_borders(t1)
    t1.autofit = False
    t1.columns[0].width = Cm(6.25)
    t1.columns[1].width = Cm(12.5)

    c = t1.cell(0, 0)
    c.width = Cm(6.25)
    p = c.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(1)
    p.add_run(department or "BỘ GIÁO DỤC VÀ ĐÀO TẠO").bold = True
    _insert_header_rule_after(c, p, cell_width_cm=6.25, width_ratio=0.70)
    p = c.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(1)
    p.add_run(exam_type or "ĐỀ THI CHÍNH THỨC")
    p = c.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(1)
    r1 = p.add_run("(Đáp án có " if is_answer else "(Đề thi có ")
    r1.italic = True
    for el in _create_field(r'NUMPAGES \# "00"', italic=True):
        p._p.append(el)
    r2 = p.add_run(" trang)")
    r2.italic = True

    c = t1.cell(0, 1)
    c.width = Cm(12.5)
    p = c.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(1)
    p.add_run(exam_title or "").bold = True
    p = c.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(1)
    p.add_run("Môn thi: " + (subject or "...")).bold = True
    p = c.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(1)
    r_dur = p.add_run(f"Thời gian làm bài: {duration or 50} phút, không kể thời gian phát đề")
    r_dur.italic = True
    _insert_header_rule_after(c, p, cell_width_cm=12.5, width_ratio=0.55)

    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(3)

    t2 = doc.add_table(rows=1, cols=2)
    _remove_borders(t2)
    t2.autofit = False
    t2.columns[0].width = Cm(12.5)
    t2.columns[1].width = Cm(6.25)

    c = t2.cell(0, 0)
    c.width = Cm(12.5)
    p = c.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    if not is_answer:
        p.add_run("Họ, tên thí sinh: ........................................................................").bold = True
        p2 = c.add_paragraph()
        p2.paragraph_format.space_after = Pt(0)
        p2.add_run("Số báo danh: .............................................................................").bold = True

    c = t2.cell(0, 1)
    c.width = Cm(6.25)
    c.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p = c.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)

    # Hộp "Mã đề" — VML `<v:rect>` cũ (đã chốt giữ nguyên như bản cũ, không
    # đổi sang ô bảng đơn giản hơn) — port nguyên xi `word_exporter.py:568-591`.
    vml_xml = f"""
    <w:r xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" xmlns:v="urn:schemas-microsoft-com:vml" xmlns:w10="urn:schemas-microsoft-com:office:word" xmlns:o="urn:schemas-microsoft-com:office:office">
      <w:pict>
        <v:rect id="_x0000_s1026" style="width:130pt;height:18pt;v-text-anchor:middle;" fillcolor="white" strokecolor="black" strokeweight="0.75pt">
          <v:textbox inset="0,0,0,0">
            <w:txbxContent>
              <w:p>
                <w:pPr>
                  <w:jc w:val="center"/>
                  <w:spacing w:after="0"/>
                </w:pPr>
                <w:r>
                  <w:rPr><w:b/></w:rPr>
                  <w:t>Mã đề: {code}</w:t>
                </w:r>
              </w:p>
            </w:txbxContent>
          </v:textbox>
        </v:rect>
      </w:pict>
    </w:r>
    """
    p._p.append(parse_xml(vml_xml))


def add_exam_footer(doc, code):
    """Chân trang "Trang X/Y - Mã đề thi {code}" — port
    `word_exporter.py:680-689`."""
    section = doc.sections[0]
    footer = section.footer
    fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    fp.add_run("Trang ")
    for el in _create_field("PAGE"):
        fp._p.append(el)
    fp.add_run("/")
    for el in _create_field("NUMPAGES"):
        fp._p.append(el)
    fp.add_run(f" - Mã đề thi {code}")


def _fill_cell(cell, text, bold=False, center=False, vcenter=False):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = bold
    if vcenter:
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def _mc_letter(q):
    for idx, opt in enumerate(q.get("options") or []):
        if opt.get("is_correct"):
            return chr(65 + idx)
    return ""


def _add_mc_answer_grid(doc, mc_list):
    """Bảng số câu/đáp án, 10 câu mỗi hàng — port cấu trúc bảng LaTeX
    `word_exporter.py:345-355`: MỘT bảng DUY NHẤT (không phải 1 bảng riêng
    mỗi lô 10 câu), nối thêm 1 cặp hàng (số câu + đáp án) cho mỗi lô 10 câu
    VÀO CÙNG bảng đó — khớp đúng bản cũ nối `\\hline` liên tục, không có
    khoảng trắng giữa các lô (bản trước của tôi tạo bảng riêng + đoạn trống
    giữa các lô, sai với bản cũ, đã sửa)."""
    items = [(i + 1, _mc_letter(q)) for i, q in enumerate(mc_list)]
    if not items:
        return
    per = 10
    ncol = min(per, len(items))
    n_chunks = (len(items) + per - 1) // per
    t = doc.add_table(rows=2 * n_chunks, cols=ncol)
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for ci, start in enumerate(range(0, len(items), per)):
        chunk = items[start:start + per]
        for j, (num, ans) in enumerate(chunk):
            _fill_cell(t.cell(2 * ci, j), str(num), center=True)
            _fill_cell(t.cell(2 * ci + 1, j), ans, center=True)
    doc.add_paragraph()


def _add_sa_answer_grid(doc, sa_list):
    """Bảng số câu/đáp án ngắn, 10 câu mỗi hàng + 1 cột nhãn đầu — port
    `word_exporter.py:373-387`, cùng nguyên tắc MỘT bảng duy nhất như
    `_add_mc_answer_grid`. Đáp án đã được ép về 1 dòng (`.replace('\\n',' ')`)
    — đây là chỗ DUY NHẤT trong toàn bộ bộ ghi tự xử lý `\\n` trên MỘT chuỗi
    thô (khác việc chuyển ngắt mềm->cứng áp dụng cho cả cây tài liệu, không
    phải riêng ô đáp án ngắn này)."""
    items = []
    for i, q in enumerate(sa_list):
        opts = q.get("options") or [{}]
        ans = str(opts[0].get("content", "")).replace("\n", " ").strip()
        items.append((i + 1, ans))
    if not items:
        return
    per = 10
    ncol = min(per, len(items))
    n_chunks = (len(items) + per - 1) // per
    t = doc.add_table(rows=2 * n_chunks, cols=ncol + 1)
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for ci, start in enumerate(range(0, len(items), per)):
        chunk = items[start:start + per]
        _fill_cell(t.cell(2 * ci, 0), "Câu", center=True)
        _fill_cell(t.cell(2 * ci + 1, 0), "Đáp án", center=True)
        for j, (num, ans) in enumerate(chunk):
            _fill_cell(t.cell(2 * ci, j + 1), str(num), center=True)
            _fill_cell(t.cell(2 * ci + 1, j + 1), ans, center=True)
    doc.add_paragraph()


def _add_tf_answer_table(doc, tf_list):
    """Bảng đáp án Đúng/Sai, gộp ô cột "Câu" — port `build_tf_table()`
    (`word_exporter.py:630-668`), chỉ đổi từ "chèn trước anchor" sang
    "append thẳng" vì dựng tuần tự, không hậu xử lý."""
    n = len(tf_list)
    if n == 0:
        return
    half = (n + 1) // 2
    table = doc.add_table(rows=1 + half * 4, cols=6)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    labels = ["a)", "b)", "c)", "d)"]

    def ds(q):
        return ["Đúng" if o.get("is_correct") else "Sai" for o in (q.get("options") or [])]

    for j, h in enumerate(["Câu", "Lệnh hỏi", "Đáp án", "Câu", "Lệnh hỏi", "Đáp án"]):
        _fill_cell(table.cell(0, j), h, center=True)

    for i in range(half):
        base = 1 + i * 4
        left = tf_list[i]
        ld = ds(left)
        ri = i + half
        right = tf_list[ri] if ri < n else None
        rd = ds(right) if right else []
        for r in range(4):
            _fill_cell(table.cell(base + r, 1), labels[r] if r < len(ld) else "", center=True)
            _fill_cell(table.cell(base + r, 2), ld[r] if r < len(ld) else "", center=True)
            _fill_cell(table.cell(base + r, 4), labels[r] if (right and r < len(rd)) else "", center=True)
            _fill_cell(table.cell(base + r, 5), rd[r] if (right and r < len(rd)) else "", center=True)
        cl = table.cell(base, 0)
        for r in range(1, 4):
            cl = cl.merge(table.cell(base + r, 0))
        _fill_cell(cl, str(i + 1), center=True, vcenter=True)
        cr = table.cell(base, 3)
        for r in range(1, 4):
            cr = cr.merge(table.cell(base + r, 3))
        if right:
            _fill_cell(cr, str(ri + 1), center=True, vcenter=True)


def add_answer_key(doc, questions, contest, code, exam_title, department, exam_type, subject, duration, math_batch=None):
    """Trang đáp án (hướng dẫn chấm) + "LỜI GIẢI CHI TIẾT" — port
    `_render_answer_key_body()` (`word_exporter.py:315-399`). Gọi SAU khi đã
    dựng xong đề sạch (thân đề `include_solution=False`) trong CÙNG 1 doc.
    """
    mc_list = [q for q in questions if q.get("question_type") == "mc"]
    tf_list = [q for q in questions if q.get("question_type") == "tf"]
    sa_list = [q for q in questions if q.get("question_type") == "sa"]
    oe_list = [q for q in questions if q.get("question_type") == "oe"]
    w = get_scoring_config(contest or {})
    has_oe = len(oe_list) > 0

    add_page_break(doc)
    add_exam_header(doc, code, exam_title, department, exam_type, subject, duration, is_answer=True)

    if has_oe:
        points_a = len(mc_list) * w["mc"] + len(tf_list) * w["tf"] + len(sa_list) * w["sa"]
        doc.add_paragraph().add_run(f"PHẦN TRẮC NGHIỆM: {fmt_points(points_a)} điểm").bold = True

    if mc_list:
        head = "I." if has_oe else "PHẦN I."
        p = doc.add_paragraph()
        p.add_run(f"{head} Câu trắc nghiệm nhiều phương án lựa chọn:").bold = True
        p.add_run(f" Mỗi câu trả lời đúng thí sinh được {fmt_points(w['mc'])} điểm.")
        _add_mc_answer_grid(doc, mc_list)

    if tf_list:
        head = "II." if has_oe else "PHẦN II."
        doc.add_paragraph().add_run(f"{head} Câu trắc nghiệm đúng sai.").bold = True
        doc.add_paragraph(f"- Thí sinh chỉ lựa chọn chính xác 01 ý trong 1 câu hỏi được {fmt_points(0.1 * w['tf'])} điểm.")
        doc.add_paragraph(f"- Thí sinh chỉ lựa chọn chính xác 02 ý trong 1 câu hỏi được {fmt_points(0.25 * w['tf'])} điểm.")
        doc.add_paragraph(f"- Thí sinh chỉ lựa chọn chính xác 03 ý trong 1 câu hỏi được {fmt_points(0.5 * w['tf'])} điểm.")
        doc.add_paragraph(f"- Thí sinh lựa chọn chính xác cả 04 ý trong 1 câu hỏi được {fmt_points(1.0 * w['tf'])} điểm.")
        _add_tf_answer_table(doc, tf_list)

    if sa_list:
        head = "III." if has_oe else "PHẦN III."
        p = doc.add_paragraph()
        p.add_run(f"{head} Câu trắc nghiệm trả lời ngắn:").bold = True
        p.add_run(f" Mỗi câu trả lời đúng thí sinh được {fmt_points(w['sa'])} điểm.")
        _add_sa_answer_grid(doc, sa_list)

    if has_oe:
        points_b = len(oe_list) * w["oe"]
        doc.add_paragraph().add_run(f"PHẦN TỰ LUẬN: {fmt_points(points_b)} điểm").bold = True

    heading = doc.add_paragraph()
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    heading.add_run("LỜI GIẢI CHI TIẾT").bold = True

    add_exam_body(doc, questions, show_answers=True, include_solution=True, math_batch=math_batch)


def render_exam_docx(
    questions, path, code="000", exam_title="", department="", exam_type="",
    subject="", duration=50, show_answers=True, include_solution=True,
    dual_section=False, contest=None, render_math=True, general_info="",
    word_equation_format="omml",
):
    """Danh sách câu (phẳng, đã JOIN `options`/`images`) -> file `.docx` hoàn
    chỉnh. `dual_section=True`: đề sạch (không lời giải) + trang đáp án +
    "LỜI GIẢI CHI TIẾT" trong CÙNG 1 file — khớp `word_exporter.py:410-415`.
    `dual_section=False`: chỉ 1 lượt thân đề theo đúng `show_answers`/
    `include_solution` truyền vào (dùng cho các mã đề đảo — thường đề sạch).
    `render_math=True` gọi coordinator 1 LẦN CUỐI (sau khi đã dựng hết cả
    file): dựng nền OMML rồi tùy chọn chuyển sang MathType — xem
    `doctree/write/docx_math.py`.
    """
    doc = setup_document(Document())
    math_batch = [] if render_math else None
    add_exam_header(doc, code, exam_title, department, exam_type, subject, duration, is_answer=False)
    if general_info:
        general_info_doc, _ = read_field(general_info)
        write_field(doc, general_info_doc, {}, math_batch=math_batch)
    if dual_section:
        add_exam_body(doc, questions, show_answers=False, include_solution=False, math_batch=math_batch)
        add_answer_key(doc, questions, contest, code, exam_title, department, exam_type, subject, duration, math_batch=math_batch)
    else:
        add_exam_body(doc, questions, show_answers=show_answers, include_solution=include_solution, math_batch=math_batch)
    add_exam_footer(doc, code)
    if render_math:
        from doctree.write.docx_math import finalize_math_document

        docx_bytes = finalize_math_document(
            doc,
            math_batch,
            equation_format=word_equation_format,
        )
    else:
        output = io.BytesIO()
        doc.save(output)
        docx_bytes = output.getvalue()
        output.close()
    if hasattr(path, "write"):
        path.write(docx_bytes)
    else:
        with open(path, "wb") as target:
            target.write(docx_bytes)
    return path


def export_word(
    contest, questions, code, exam_title, department, exam_type, subject,
    duration, general_info, include_solution, zf, dual_section=False,
    show_answers=None, word_equation_format="omml",
):
    """CÙNG chữ ký với `word_exporter.py:802::export_word` — để
    `export_manager.py` chỉ cần đổi 1 dòng import (`from .word_exporter
    import export_word` -> `from .docx_exporter import export_word`) khi tới
    lúc nối thật (Phase 9), không đụng gì tới logic shuffle/zip/Excel đáp án
    của `export_manager.py`.

    `general_info` là LaTeX được sinh từ TreeDoc ở modal xuất đề; đọc lại qua
    doctree và chèn ngay sau header để PDF/LaTeX/Word cùng một nội dung.
    `show_answers=None` -> mặc định bằng `include_solution`, khớp mọi lời
    gọi hiện có của bản cũ (bản cũ chỉ có 1 cờ `include_solution`).
    """
    if show_answers is None:
        show_answers = include_solution
    buf = io.BytesIO()
    try:
        render_exam_docx(
            questions, buf, code=code, exam_title=exam_title, department=department,
            exam_type=exam_type, subject=subject, duration=duration,
            show_answers=show_answers, include_solution=include_solution,
            dual_section=dual_section, contest=contest, render_math=True,
            general_info=general_info, word_equation_format=word_equation_format,
        )
        # Không dùng `buf.getvalue()`: nó tạo thêm một bản sao nguyên DOCX
        # trong RAM. Stream thẳng từng khối vào ZIP rồi đóng buffer ngay.
        buf.seek(0)
        with zf.open(f"{code}.docx", "w") as target:
            shutil.copyfileobj(buf, target, length=1024 * 1024)
    finally:
        buf.close()
        # python-docx/lxml có các vòng tham chiếu; thu gom sau từng mã để 24
        # đề không dồn object XML của các mã trước tới cuối toàn bộ job.
        gc.collect()
