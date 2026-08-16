"""Bộ ghi: cây tài liệu -> `.docx`, dựng thẳng bằng `python-docx`.

Ảnh Word được dựng inline để có bố cục ổn định giữa các phiên bản Word. Với
 cây `side:left/right`, ảnh được đặt sau nội dung hỏi / trước các phương án
thay vì đổi thành `<wp:anchor>`: anchor ngoài luồng có thể làm vùng wrap lấn
lên câu trước dù XML neo đúng paragraph. Ảnh không có side giữ vị trí trong cây.
"""
import atexit
import contextvars
import hashlib
import os
import re
import shutil
import subprocess
import tempfile
from functools import wraps

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import (
    WD_CELL_VERTICAL_ALIGNMENT,
    WD_ROW_HEIGHT_RULE,
    WD_TABLE_ALIGNMENT,
)
from docx.oxml import OxmlElement
from docx.oxml import parse_xml
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.opc.part import Part
from docx.shared import Inches, Pt, Cm, RGBColor


def _apply_paragraph_alignment(paragraph, node):
    paragraph.alignment = {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
        "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
    }.get(node.get("align", "left"), WD_ALIGN_PARAGRAPH.LEFT)

from ..figures import A4_REFERENCE_WIDTH_IN, svg_native_width_in, resolve_image_file
from ..adapt import content_len, has_image

TEXT_WIDTH_IN = 6.5      # vùng chữ; chỉ dùng cho cột/minipage
IMAGE_REFERENCE_WIDTH_IN = A4_REFERENCE_WIDTH_IN  # 8.2677in = toàn bộ A4

# Hệ tọa độ ngang của đáp án Word (cm). Tab stop trong Word được tính từ mép
# trái vùng chữ của trang, trong khi ký tự A bắt đầu sau `left_indent`. Vì vậy
# phải chia phần chiều rộng CÒN LẠI sau indent, không dùng lại 1/4 của toàn bộ
# vùng chữ (4.5625 cm) như trước.
DOCX_A4_WIDTH_CM = 21.0
DOCX_LEFT_MARGIN_CM = 1.5
DOCX_RIGHT_MARGIN_CM = 1.2
DOCX_OPTION_LEFT_INDENT_CM = 0.5
DOCX_TEXT_WIDTH_CM = DOCX_A4_WIDTH_CM - DOCX_LEFT_MARGIN_CM - DOCX_RIGHT_MARGIN_CM
DOCX_OPTION_WIDTH_CM = DOCX_TEXT_WIDTH_CM - DOCX_OPTION_LEFT_INDENT_CM


def _option_tab_positions_cm(columns):
    """Vị trí tuyệt đối của các cột 2..N trong vùng chữ Word."""
    column_width = DOCX_OPTION_WIDTH_CM / columns
    return tuple(
        DOCX_OPTION_LEFT_INDENT_CM + column_width * index
        for index in range(1, columns)
    )

_SVG_FALLBACK_DIR = contextvars.ContextVar("docx_svg_fallback_dir", default=None)
_STANDALONE_FALLBACK_DIR = None


def with_svg_fallback_cache(func):
    """Một cache PNG trên đĩa cho toàn bộ lượt xuất ZIP, tự xóa khi xong.

    Nhờ scope này, cùng một SVG xuất hiện trong đề gốc và 24 mã đảo chỉ được
    rasterize đúng một lần; cache chỉ giữ đường dẫn, không giữ PNG trong RAM.
    """
    @wraps(func)
    def wrapped(*args, **kwargs):
        with tempfile.TemporaryDirectory(prefix="docx-svg-fallback-") as cache_dir:
            token = _SVG_FALLBACK_DIR.set(cache_dir)
            try:
                return func(*args, **kwargs)
            finally:
                _SVG_FALLBACK_DIR.reset(token)
    return wrapped


def _resolve_path(row):
    """`storage_path` lưu trong CSDL là URL web (vd `/static/images/2026/08/
    xxx.svg`), KHÔNG phải đường dẫn file thật trên đĩa — phải resolve qua
    `doctree.figures.resolve_image_file` trước khi `add_picture()`/đọc file
    SVG, nếu không sẽ `FileNotFoundError` ngay khi xuất (đã gặp thật khi
    người dùng xuất đề qua web)."""
    path = row.get("storage_path") or ""
    if not path:
        return path
    resolved, _matched = resolve_image_file(path, [row])
    return resolved


def _standalone_fallback_dir():
    """Cache dự phòng cho các lời gọi render DOCX lẻ, ngoài luồng xuất ZIP."""
    global _STANDALONE_FALLBACK_DIR
    if _STANDALONE_FALLBACK_DIR is None:
        _STANDALONE_FALLBACK_DIR = tempfile.mkdtemp(prefix="docx-svg-standalone-")
        atexit.register(shutil.rmtree, _STANDALONE_FALLBACK_DIR, ignore_errors=True)
    return _STANDALONE_FALLBACK_DIR


def _svg_fallback_path(svg_path):
    """Tạo PNG fallback 300 DPI trên đĩa; không giữ bytes ảnh trong RAM."""
    stat = os.stat(svg_path)
    cache_key = hashlib.sha256(
        f"{os.path.abspath(svg_path)}|{stat.st_mtime_ns}|{stat.st_size}".encode()
    ).hexdigest()
    cache_dir = _SVG_FALLBACK_DIR.get() or _standalone_fallback_dir()
    png_path = os.path.join(cache_dir, f"{cache_key}.png")
    if os.path.isfile(png_path):
        return png_path

    magick = shutil.which("magick")
    if not magick:
        raise RuntimeError(
            f"Không thể chèn SVG vào DOCX vì thiếu ImageMagick: {svg_path}"
        )
    subprocess.run(
        [magick, "-density", "300", svg_path, png_path],
        check=True, capture_output=True, timeout=60,
    )
    return png_path


def _resolve_picture_source(resolved_path):
    """Nguồn cho ``python-docx.add_picture()``.

    Hàm cấp cao đó không đọc SVG, nên SVG dùng PNG 300 DPI trên đĩa làm
    ``a:blip``. Ngay sau đó `_attach_svg_to_picture()` gắn SVG nguyên bản
    bằng ``asvg:svgBlip`` để Word mới vẫn hiển thị vector.
    """
    if not resolved_path.lower().endswith(".svg"):
        return resolved_path
    return _svg_fallback_path(resolved_path)


def _attach_svg_to_picture(run, svg_path):
    """Gắn SVG vector vào ảnh PNG fallback vừa tạo bằng DrawingML Office.

    Word mới ưu tiên ``asvg:svgBlip``; phần ``a:blip`` do python-docx tạo
    vẫn là PNG fallback cho Word cũ/LibreOffice/trình xem web.
    """
    with open(svg_path, "rb") as source:
        svg_blob = source.read()
    svg_key = hashlib.sha256(svg_blob).hexdigest()

    # python-docx tự khử trùng PNG theo SHA1; Part SVG tự tạo cũng cần làm
    # tương tự để một hình lặp ở đề + lời giải không phình DOCX gấp đôi.
    svg_parts = getattr(run.part, "_exam_svg_parts", None)
    if svg_parts is None:
        svg_parts = {}
        run.part._exam_svg_parts = svg_parts
    svg_part = svg_parts.get(svg_key)
    if svg_part is None:
        package = run.part.package
        partname = package.next_partname("/word/media/image%d.svg")
        svg_part = Part(partname, "image/svg+xml", svg_blob, package)
        svg_parts[svg_key] = svg_part
    svg_rid = run.part.relate_to(svg_part, RT.IMAGE)

    blips = run._r.xpath(".//a:blip")
    if not blips:
        raise RuntimeError("Không tìm thấy a:blip để gắn SVG vào ảnh DOCX")
    blip = blips[0]
    ext_lst = blip.find(qn("a:extLst"))
    if ext_lst is None:
        ext_lst = OxmlElement("a:extLst")
        blip.append(ext_lst)
    ext_lst.append(parse_xml(
        '<a:ext xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main" '
        'xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships" '
        'xmlns:asvg="http://schemas.microsoft.com/office/drawing/2016/SVG/main" '
        'uri="{96DAC541-7B7A-43D3-8B79-37D633B846F1}">'
        f'<asvg:svgBlip r:embed="{svg_rid}"/>'
        '</a:ext>'
    ))


def _add_picture(run, resolved_path, width=None):
    """Chèn ảnh; riêng SVG dùng PNG fallback + SVG vector nguyên bản."""
    picture_source = _resolve_picture_source(resolved_path)
    if width is not None:
        run.add_picture(picture_source, width=width)
    else:
        run.add_picture(picture_source)
    if resolved_path.lower().endswith(".svg"):
        _attach_svg_to_picture(run, resolved_path)
    return run


def _pic_width_in(row, resolved_path=None):
    """`width` (tỉ lệ 0–1) -> số inch để truyền cho `add_picture()`.

    `None` — chưa đặt tỉ lệ — trả lại `None` luôn: với ảnh raster (PNG/JPG),
    không truyền `width=` thì `python-docx` tự đọc kích thước gốc qua PIL.
    Ảnh `.svg` (TikZ) thì PIL không mở được, nên phải tự đọc `width`/`viewBox`
    của chính file SVG (`svg_native_width_in`) rồi truyền `Inches(...)` tường
    minh — không có cách nào bỏ trống cho loại này.
    """
    w = row.get("width")
    if w is not None:
        return Inches(IMAGE_REFERENCE_WIDTH_IN * float(w))
    path = resolved_path if resolved_path is not None else _resolve_path(row)
    if path.lower().endswith(".svg"):
        return Inches(svg_native_width_in(path))
    return None


# Ảnh neo, chữ trôi

def float_picture(run, side="right", dist_pt=8):
    """Đổi ảnh inline trong `run` thành ảnh NEO có chữ trôi quanh."""
    drawing = run._r.find(qn("w:drawing"))
    inline = drawing[0]
    extent = inline.find(qn("wp:extent"))
    doc_pr = inline.find(qn("wp:docPr"))
    frame_pr = inline.find(qn("wp:cNvGraphicFramePr"))
    graphic = inline.find(qn("a:graphic"))

    emu = int(dist_pt * 12700)
    anchor = OxmlElement("wp:anchor")
    for k, v in {"distT": "0", "distB": "0", "distL": str(emu), "distR": str(emu),
                 "simplePos": "0", "relativeHeight": "2", "behindDoc": "0",
                 "locked": "0", "layoutInCell": "1", "allowOverlap": "1"}.items():
        anchor.set(k, v)

    sp = OxmlElement("wp:simplePos")
    sp.set("x", "0")
    sp.set("y", "0")
    anchor.append(sp)

    ph = OxmlElement("wp:positionH")
    ph.set("relativeFrom", "margin")
    al = OxmlElement("wp:align")
    al.text = side
    ph.append(al)
    anchor.append(ph)

    pv = OxmlElement("wp:positionV")
    pv.set("relativeFrom", "paragraph")
    off = OxmlElement("wp:posOffset")
    off.text = "0"
    pv.append(off)
    anchor.append(pv)

    anchor.append(extent)

    ee = OxmlElement("wp:effectExtent")
    for k in "ltrb":
        ee.set(k, "0")
    anchor.append(ee)

    wrap = OxmlElement("wp:wrapSquare")          # <- thứ làm chữ trôi quanh ảnh
    wrap.set("wrapText", "bothSides")
    anchor.append(wrap)

    anchor.append(doc_pr)
    if frame_pr is not None:
        anchor.append(frame_pr)
    anchor.append(graphic)

    drawing.remove(inline)
    drawing.append(anchor)


# Dựng nội dung

def add_inline(par, nodes, figures, math_batch=None, host=None):
    """`math_batch` — nếu truyền một list, node `math` sẽ được GHI VÀO đó
    dưới dạng `("inline", run, tex)` (bên cạnh việc vẫn dựng placeholder chữ
    nghiêng TeX thô như trước) để lượt 2 (`docx_math` điều phối sang
    `docx_math_omml`) thay bằng công thức OMML thật. Không truyền
    `math_batch` (None) thì giữ nguyên hành vi cũ —
    chỉ chữ nghiêng, dùng cho test/gọi lẻ không cần công thức thật.

    `host` — nơi chứa `par` (`doc`/ô bảng), có thể gọi `.add_paragraph()`.
    Nếu truyền: gặp `hard_break` (Shift+Enter trong cây) sẽ TÁCH hẳn thành
    MỘT ĐOẠN VĂN MỚI (`<w:p>`, Enter thật) thay vì chỉ xuống dòng trong CÙNG
    một đoạn (`<w:br/>`, Shift+Enter) — đúng yêu cầu người dùng. Lý do:
    `<w:br/>` bên trong một đoạn canh Justify (mặc định `setup_document`) bị
    Word tự kéo giãn từng dòng (trừ dòng cuối) cho đầy hết chiều ngang —
    thấy rõ ở phần "Lời giải" nhiều câu (vd câu có `hard_break` giữa các ý),
    chữ bị dãn ra trông rất xấu. Tách đoạn thật thì mỗi dòng là 1 đoạn ĐỘC
    LẬP, không còn bị dãn theo kiểu đó. KHÔNG truyền `host` (None) thì giữ
    `add_break()` cũ (dùng ở nơi không có quyền tạo đoạn mới, vd nội dung 1
    ô bảng ngắn không cần tách). Trả về đoạn ĐANG DÙNG sau cùng, để caller
    nối tiếp nếu còn nội dung khác sau lời gọi này (xem `write_field`).
    """
    for n in nodes:
        t = n["type"]
        if t == "text":
            marks = n.get("marks", [])
            # Cùng logic hard_break bên dưới — nếu chuỗi lỡ có `\n` thô (vd
            # dữ liệu import cũ) thì coi như một lần ngắt mềm nữa, xử lý
            # nhất quán (tách đoạn nếu có host, không thì `<w:br/>`).
            lines = n["text"].split("\n")
            for i, line in enumerate(lines):
                if i > 0:
                    if host is not None:
                        par = host.add_paragraph()
                    else:
                        par.add_run().add_break()
                r = par.add_run(line)
                r.bold = "bold" in marks
                r.italic = "italic" in marks
                r.underline = "underline" in marks
                color = n.get("color")
                if color and re.fullmatch(r"#[0-9A-Fa-f]{6}", color):
                    r.font.color.rgb = RGBColor.from_string(color[1:].upper())
        elif t == "math":
            # Placeholder chữ nghiêng TeX thô — lượt 2 thay bằng OMML thật
            # nếu `math_batch` được truyền; là bản hiển thị CUỐI CÙNG nếu
            # không (hoặc nếu pandoc lỗi ở lượt 2, xem `render_exam_docx`).
            r = par.add_run(n["tex"])
            r.italic = True
            r.font.name = "Cambria Math"
            if math_batch is not None:
                math_batch.append(("inline", r, n["tex"]))
        elif t == "hard_break":
            if host is not None:
                par = host.add_paragraph()
            else:
                par.add_run().add_break()
        elif t == "image_inline":
            row = (figures or {}).get(n["figure_id"])
            if row and row.get("storage_path"):
                path = _resolve_path(row)
                w = _pic_width_in(row, resolved_path=path)
                _add_picture(par.add_run(), path, width=w)


def add_blocks(host, nodes, figures, math_batch=None):
    for n in nodes:
        t = n["type"]
        if t == "paragraph":
            paragraph = host.add_paragraph()
            _apply_paragraph_alignment(paragraph, n)
            add_inline(paragraph, n["content"], figures, math_batch, host=host)
        elif t == "math_block":
            p = host.add_paragraph()
            # Căn TRÁI (không phải giữa) — khớp hành vi ép-trái công thức
            # khối của bộ cũ (`word_exporter.py:691-732`, oMathPara mặc định
            # Word/pandoc căn giữa nhưng bộ cũ luôn ép về trái cho khớp bố
            # cục đề thi). Writer `docx_math_omml` áp lại đúng việc này lên
            # chính OMML thật ở lượt 2.
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.add_run(n["tex"]).italic = True
            if math_batch is not None:
                math_batch.append(("display", p, n["tex"]))
        elif t == "image":
            p = host.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_inline(p, [{"type": "image_inline", "figure_id": n["figure_id"]}], figures, math_batch)
        elif t == "table":
            add_table(host, n, figures, math_batch)
        elif t == "list":
            add_list(host, n, figures, math_batch)
        elif t == "columns":
            add_columns(host, n, figures, math_batch)


def add_columns(host, node, figures, math_batch=None):
    """Dựng columns bằng bảng bố cục không viền; cho phép bảng lồng trong ô."""
    columns = node.get("columns", [])
    if not columns:
        return
    table = host.add_table(rows=1, cols=len(columns))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    tbl_borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border = OxmlElement(f"w:{edge}")
        border.set(qn("w:val"), "nil")
        tbl_borders.append(border)
    table._tbl.tblPr.append(tbl_borders)

    align_of = {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
    }
    for cell, column in zip(table.rows[0].cells, columns):
        width = Inches(TEXT_WIDTH_IN * float(column["width"]))
        cell.width = width
        tc_w = cell._tc.get_or_add_tcPr().get_or_add_tcW()
        tc_w.set(qn("w:w"), str(int(width.twips)))
        tc_w.set(qn("w:type"), "dxa")
        cell.vertical_alignment = {
            "top": WD_CELL_VERTICAL_ALIGNMENT.TOP,
            "center": WD_CELL_VERTICAL_ALIGNMENT.CENTER,
            "bottom": WD_CELL_VERTICAL_ALIGNMENT.BOTTOM,
        }.get(column.get("valign", "top"), WD_CELL_VERTICAL_ALIGNMENT.TOP)
        # Paragraph bắt buộc sẵn có trong cell: giữ nhưng triệt khoảng cách;
        # nội dung block (kể cả bảng dữ liệu) được thêm ngay sau nó.
        first = cell.paragraphs[0]
        first.alignment = align_of.get(column.get("align", "left"))
        first.paragraph_format.space_before = Pt(0)
        first.paragraph_format.space_after = Pt(0)
        add_blocks(cell, column.get("content", []), figures, math_batch)
        if column.get("content") and first._p.getparent() is not None:
            first._p.getparent().remove(first._p)
        for paragraph in cell.paragraphs:
            paragraph.alignment = align_of.get(column.get("align", "left"))


def add_list(host, node, figures, math_batch=None, level=0):
    """Dựng list Word đệ quy; không bỏ mất list/table nằm trong một item.

    Word có sẵn style `List Bullet/Number`, `... 2`, `... 3`. Cấp sâu hơn 3
    dùng lại style cấp 3 thay vì làm hỏng file vì tham chiếu style không có.
    Paragraph đầu của item mang bullet/number; các block sau thuộc cùng item
    được dựng tiếp, list con tăng đúng một cấp.
    """
    base = "List Number" if node.get("ordered") else "List Bullet"
    style_level = min(level + 1, 3)
    style = base if style_level == 1 else f"{base} {style_level}"
    for item in node.get("items", []):
        first_paragraph = True
        for block in item:
            kind = block.get("type")
            if kind == "paragraph":
                if first_paragraph:
                    p = host.add_paragraph(style=style)
                    first_paragraph = False
                else:
                    p = host.add_paragraph()
                    p.paragraph_format.left_indent = Inches(0.25 * (level + 1))
                add_inline(p, block["content"], figures, math_batch, host=host)
            elif kind == "list":
                add_list(host, block, figures, math_batch, level=level + 1)
            else:
                # Bảng, ảnh, công thức khối... vẫn phải được giữ; add_blocks
                # xử lý đúng từng loại thay vì nhánh list cũ âm thầm bỏ qua.
                add_blocks(host, [block], figures, math_batch)


def add_table(host, node, figures, math_batch=None):
    occupied = {}
    placements = []
    ncol = 0
    for r_idx, row_data in enumerate(node["rows"]):
        cursor = 0
        for cell_data in row_data:
            while (r_idx, cursor) in occupied:
                cursor += 1
            colspan = int(cell_data.get("colspan", 1))
            rowspan = int(cell_data.get("rowspan", 1))
            placements.append((r_idx, cursor, cell_data, rowspan, colspan))
            for rr in range(r_idx, r_idx + rowspan):
                for cc in range(cursor, cursor + colspan):
                    occupied[(rr, cc)] = True
            cursor += colspan
        ncol = max(ncol, cursor, max((c + 1 for (r, c) in occupied if r == r_idx), default=0))
    t = host.add_table(rows=len(node["rows"]), cols=ncol)
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False
    # Full bề ngang trang (pct=5000 == 100.00% theo đơn vị OOXML) — đặt lúc
    # tạo thay vì quét lại mọi bảng sau khi dựng xong như bộ cũ
    # (`word_exporter.py:758-772`). Style "Table Grid" đã tự chèn sẵn 1
    # `w:tblW` (auto/0) — phải GỠ nó trước rồi chèn cái mới ĐÚNG VỊ TRÍ theo
    # thứ tự bắt buộc của schema (`tblW` phải đứng TRƯỚC `jc`, `t.alignment`
    # ở trên đã tạo `jc`) — nếu chỉ `append()` thẳng sẽ ra 2 `tblW` trùng và
    # sai thứ tự, Word có thể báo cần sửa file.
    tbl_pr = t._tbl.tblPr
    for old_w in tbl_pr.findall(qn("w:tblW")):
        tbl_pr.remove(old_w)
    tbl_w = OxmlElement("w:tblW")
    tbl_w.set(qn("w:w"), "5000")
    tbl_w.set(qn("w:type"), "pct")
    jc = tbl_pr.find(qn("w:jc"))
    if jc is not None:
        jc.addprevious(tbl_w)
    else:
        tbl_pr.append(tbl_w)

    # Bảng vẫn chiếm toàn bộ vùng chữ, nhưng cột không bị chia đều máy móc.
    # Lấy độ dài nội dung lớn nhất của từng cột làm trọng số để cột ngắn chỉ
    # nhận phần vừa đủ, phần còn lại dành cho cột có nội dung dài hơn.
    def cell_weight(cell):
        score = 0
        for item in cell.get("content", []):
            kind = item.get("type")
            if kind == "text":
                score += len(item.get("text", ""))
            elif kind == "math":
                score += max(3, len(item.get("tex", "")) // 2)
            elif kind == "image_inline":
                score += 12
            else:
                score += 4
        return max(3, score + 2)

    weights = [3] * ncol
    for _, col, cell, _, colspan in placements:
        share = max(3, cell_weight(cell) // colspan)
        for index in range(col, min(col + colspan, ncol)):
            weights[index] = max(weights[index], share)
    table_width_in = DOCX_TEXT_WIDTH_CM / 2.54
    configured_widths = node.get("widths") or []
    if len(configured_widths) == ncol and sum(configured_widths) > 0:
        configured_total = sum(configured_widths)
        widths = [
            Inches(table_width_in * float(width) / configured_total)
            for width in configured_widths
        ]
    else:
        total_weight = sum(weights) or ncol
        widths = [Inches(table_width_in * weight / total_weight) for weight in weights]
    for grid_col, width in zip(t._tbl.tblGrid.gridCol_lst, widths):
        grid_col.set(qn("w:w"), str(int(width.twips)))

    for row in t.rows:
        for i, cell in enumerate(row.cells):
            cell.width = widths[i]
            tc_w = cell._tc.get_or_add_tcPr().get_or_add_tcW()
            tc_w.set(qn("w:w"), str(int(widths[i].twips)))
            tc_w.set(qn("w:type"), "dxa")
    for index, height in enumerate(node.get("row_heights") or []):
        if index >= len(t.rows):
            break
        t.rows[index].height = Pt(float(height) * 0.75)
        t.rows[index].height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
    for r_idx, col, cell_data, rowspan, colspan in placements:
        cell = t.cell(r_idx, col)
        if rowspan > 1 or colspan > 1:
            cell = cell.merge(t.cell(r_idx + rowspan - 1, col + colspan - 1))
        cell.text = ""
        add_inline(cell.paragraphs[0], cell_data["content"], figures, math_batch, host=cell)


def write_field(doc, tree, figures, lead=None, math_batch=None,
                defer_side_images=False, lead_indent_cm=None):
    """Ghi ruột một trường.

    Ảnh của cây `side:left/right` không còn dựng thành `wp:anchor`: lấy chúng
    khỏi vị trí đầu cây và dựng inline căn giữa ở cuối trường. Riêng đề bài
    của một câu truyền `defer_side_images=True` để caller chèn ảnh ngay SAU
    nội dung hỏi nhưng TRƯỚC các phương án. Ảnh của cây không có side vẫn đi qua
    `add_blocks` đúng vị trí vốn có trong cây.

    Trả về danh sách image node đang hoãn (rỗng nếu đã dựng tại chỗ)."""
    side = tree.get("side", "center")
    floating = side in ("left", "right")
    blocks = [n for n in tree.get("content", [])
              if not (floating and n["type"] == "image")]
    figs = [n for n in tree.get("content", []) if n["type"] == "image"] if floating else []

    first = None
    if (blocks and blocks[0]["type"] == "paragraph") or lead:
        first = doc.add_paragraph()

    if first is not None:
        if lead_indent_cm is not None:
            first.paragraph_format.left_indent = Cm(lead_indent_cm)
        if lead:
            first.add_run(lead).bold = True
        if blocks and blocks[0]["type"] == "paragraph":
            _apply_paragraph_alignment(first, blocks[0])
            add_inline(first, blocks[0]["content"], figures, math_batch, host=doc)
            blocks = blocks[1:]

    add_blocks(doc, blocks, figures, math_batch)
    if figs and not defer_side_images:
        add_blocks(doc, figs, figures, math_batch)
        return []
    return figs


def _write_option_body(p, opt, figures, math_batch=None, host=None):
    for b in opt["content_doc"]["content"]:
        if b["type"] == "paragraph":
            _apply_paragraph_alignment(p, b)
            add_inline(p, b["content"], figures, math_batch, host=host)
        elif b["type"] == "image":
            # Trong phương án, ảnh block vẫn phải nằm NGAY SAU nhãn A/B/C/D
            # để layout 2/4 cột ra dạng:
            #   A. [ảnh] <TAB> B. [ảnh]
            # Không dùng add_blocks() vì nó sẽ tạo paragraph mới và phá hàng
            # tab của các phương án đang nằm cạnh nhau.
            add_inline(
                p,
                [{"type": "image_inline", "figure_id": b["figure_id"]}],
                figures,
                math_batch,
                host=None,
            )


def _add_mc_options(doc, opts, layout_type, figures, show_answers=False, math_batch=None):
    """A/B/C/D — tự chọn 1/2/4 phương án mỗi dòng theo độ dài/có ảnh hay
    không (hoặc ép qua `layout_type`), dùng TAB STOP THẬT của python-docx —
    không cần marker `[TAB2]`/`[TAB4]` như bộ cũ (bộ cũ cần marker để sống
    sót qua pandoc; ở đây dựng trực tiếp bằng python-docx nên không cần).
    Ngưỡng độ dài khớp `word_exporter.py:227-239`.
    """
    if not opts:
        return
    has_img = any(has_image(o["content_doc"]) for o in opts)
    max_len = max((content_len(o["content_doc"]) for o in opts), default=0)

    lt = str(layout_type or "")
    if lt == "1":
        per_line = 1
    elif lt == "2":
        per_line = 2
    elif lt == "4":
        per_line = 4
    elif has_img or max_len > 35:
        per_line = 1
    elif max_len > 12:
        per_line = 2
    else:
        per_line = 4

    def add_one(p, idx, host=None):
        label = p.add_run(chr(65 + idx))
        label.bold = True
        if opts[idx].get("is_correct") and show_answers:
            label.underline = True
            label.font.color.rgb = RGBColor(0xD3, 0x2F, 0x2F)
        punctuation = p.add_run(". ")
        punctuation.bold = True
        _write_option_body(p, opts[idx], figures, math_batch, host=host)

    if per_line == 1:
        # Mỗi phương án 1 đoạn riêng — an toàn để tách thêm đoạn nếu có
        # hard_break bên trong (host=doc).
        for i in range(len(opts)):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(DOCX_OPTION_LEFT_INDENT_CM)
            add_one(p, i, host=doc)
        return

    # 2/4 phương án CHUNG 1 đoạn, ngăn cách bằng tab stop — hard_break bên
    # trong 1 phương án ở đây KHÔNG tách đoạn được (sẽ phá layout cột theo
    # tab), nên KHÔNG truyền host, giữ `add_break()` cũ cho trường hợp hiếm
    # gặp này (phương án ngắn thường không có hard_break).
    tab_positions = tuple(Cm(pos) for pos in _option_tab_positions_cm(per_line))
    for start in range(0, len(opts), per_line):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(DOCX_OPTION_LEFT_INDENT_CM)
        for pos in tab_positions:
            p.paragraph_format.tab_stops.add_tab_stop(pos)
        group = range(start, min(start + per_line, len(opts)))
        for j, idx in enumerate(group):
            if j > 0:
                p.add_run("\t")
            add_one(p, idx)


def _add_tf_options(doc, opts, figures, show_answers=False, math_batch=None):
    """a)/b)/c)/d) — luôn mỗi phương án một dòng (không gộp tab như mc)."""
    for i, o in enumerate(opts):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.5)
        label = p.add_run(chr(97 + i))
        label.bold = True
        if o.get("is_correct") and show_answers:
            label.underline = True
            label.font.color.rgb = RGBColor(0xD3, 0x2F, 0x2F)
        punctuation = p.add_run(") ")
        punctuation.bold = True
        _write_option_body(p, o, figures, math_batch, host=doc)


def add_question(doc, rec, num, figures=None, show_answers=True, include_solution=True, math_batch=None):
    """Một câu -> các đoạn trong tài liệu Word.

    `include_solution=False` bỏ hẳn lời giải (đề sạch cho học sinh).
    `show_answers=False` (mà `include_solution=True`, vd trang lời giải cho
    giáo viên nhưng chưa muốn lộ đáp án) chỉ bỏ dòng "Chọn X"/tô đúng-sai,
    KHÔNG bỏ lời giải/giải thích — khớp pattern đã dùng ở `write/tex.py`.
    `math_batch` — xem `add_inline`.
    """
    qt = rec["question_type"]
    opts = rec.get("options") or []
    side_images = write_field(
        doc, rec["content_doc"], figures, lead=f"Câu {num}: ",
        math_batch=math_batch, defer_side_images=True,
    )

    # Ảnh từng được khai báo side:left/right: đặt ngay dưới nội dung hỏi và
    # trước các phương án A/B/C/D hoặc a/b/c/d.
    if side_images:
        add_blocks(doc, side_images, figures, math_batch)

    if qt == "mc":
        _add_mc_options(
            doc, opts,
            rec.get("word_option_layout") or rec.get("layout_type"),
            figures, show_answers, math_batch,
        )
    elif qt == "tf":
        _add_tf_options(doc, opts, figures, show_answers, math_batch)

    sol = rec.get("solution_doc") or {}

    def add_solution_heading():
        # "Lời giải" là đoạn RIÊNG, CĂN GIỮA, in đậm — khớp nguyên xi
        # `word_exporter.py` (`[CENTER]\textbf{Lời giải}`, LUÔN in ra khi
        # include_solution=True, kể cả khi bản thân nội dung lời giải rỗng,
        # vì luôn có "Chọn X"/"Trả lời ngắn"/giải thích Đúng-Sai theo sau).
        # KHÔNG dính liền vào đoạn nội dung như một "lead" — khác thiết kế
        # ban đầu, sửa lại cho đúng thị giác bản cũ.
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run("Lời giải").bold = True

    if qt == "mc":
        if include_solution:
            add_solution_heading()
            if content_len(sol) > 0:
                write_field(doc, sol, figures, math_batch=math_batch)
            correct = next((o for o in opts if o["is_correct"]), None)
            if show_answers and correct is not None:
                label = chr(65 + opts.index(correct))
                doc.add_paragraph().add_run(f"Chọn {label}").bold = True
    elif qt == "tf":
        if include_solution:
            add_solution_heading()
            if content_len(sol) > 0:
                write_field(doc, sol, figures, math_batch=math_batch)
            for i, o in enumerate(opts):
                expl = o.get("explaination_doc")
                if not expl or content_len(expl) <= 0:
                    continue
                status = "Đúng" if (show_answers and o["is_correct"]) else \
                    ("Sai" if show_answers else "")
                lead = f"{chr(97 + i)}) {status}. " if status else f"{chr(97 + i)}) "
                # Dựng toàn bộ TreeDoc: đoạn chữ đầu nối ngay sau nhãn;
                # bảng/ảnh/list/cột phía sau được add_blocks tạo đúng loại
                # phần tử thay vì bị bỏ qua như vòng lặp chỉ nhận paragraph.
                write_field(
                    doc, expl, figures, lead=lead, math_batch=math_batch,
                )
    elif qt == "sa":
        # Không có "phương án" để hiện trên đề sạch — chỉ hiện ở lời giải,
        # khớp `word_exporter.py:269-276` (câu trả lời ngắn không có gì để
        # in ngoài đề bài khi include_solution=False).
        if include_solution and opts:
            add_solution_heading()
            p = doc.add_paragraph()
            p.add_run("Trả lời ngắn: ").bold = True
            if show_answers:
                _write_option_body(p, opts[0], figures, math_batch, host=doc)
            if content_len(sol) > 0:
                write_field(doc, sol, figures, math_batch=math_batch)

def setup_document(doc):
    """Khổ trang/lề/font mặc định — đặt 1 LẦN lúc tạo tài liệu, thay vì dò
    lại từng đoạn sau khi dựng xong như bộ cũ (`word_exporter.py` phải làm
    vậy vì đi qua pandoc, không kiểm soát được lúc dựng). Khổ A4, lề khớp
    `scripts/create_ref.py` (reference.docx của bộ cũ) để hai bản ra cùng
    kích thước trang dù không dùng chung cơ chế reference-doc nào.
    """
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.2)
    section.top_margin = Cm(1.15)
    section.bottom_margin = Cm(1.2)

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return doc


def to_docx(recs, path, figures=None, title=None, show_answers=True, include_solution=True, render_math=True):
    """Danh sách câu -> file `.docx`. `render_math=True` gọi pandoc 1 lần
    cuối để thay công thức chữ nghiêng bằng OMML thật (`docx_math` điều phối
    sang `docx_math_omml`);
    `False` giữ nguyên chữ nghiêng (test nhanh không cần pandoc)."""
    doc = setup_document(Document())
    if title:
        doc.add_heading(title, level=1)
    math_batch = [] if render_math else None
    for i, rec in enumerate(recs, 1):
        add_question(doc, rec, i, figures, show_answers=show_answers, include_solution=include_solution, math_batch=math_batch)
    if render_math:
        from .docx_math import flush_math_batch
        flush_math_batch(doc, math_batch)
    doc.save(path)
    return path
