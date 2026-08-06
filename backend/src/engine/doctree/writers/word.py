"""Bộ ghi: cây tài liệu -> `.docx`, dựng thẳng bằng `python-docx`.

Vì sao không qua pandoc. Đường cũ (`word_exporter.py:84`) biến mọi ảnh thành
`[CENTER]\\includegraphics{...}`, tức luôn là một đoạn riêng căn giữa — không có
nhánh nào đọc `layout_type`, và pandoc cũng không diễn tả được ảnh neo có chữ
trôi quanh. Dù dữ liệu ghi `immini` thì bản Word vẫn ra ảnh nằm giữa.

Chỗ mấu chốt ở `float_picture()`: `python-docx` chỉ dựng được ảnh *inline*, muốn
ảnh neo thì phải đổi thẻ `<wp:inline>` thành `<wp:anchor>` rồi chèn
`<wp:wrapSquare>`. Thứ tự các thẻ con là **bắt buộc** theo lược đồ OOXML — sai
thứ tự thì Word báo file hỏng.
"""
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

TEXT_WIDTH_IN = 6.5      # A4 trừ lề — dùng để quy `width` (tỉ lệ 0–1) ra inch


# ------------------------------------------------------ ảnh neo, chữ trôi ---

def float_picture(run, side="right", dist_pt=8):
    """Đổi ảnh inline trong `run` thành ảnh NEO có chữ trôi quanh."""
    drawing = run._r.find(qn("w:drawing"))
    inline = drawing[0]
    extent = inline.find(qn("wp:extent"))
    doc_pr = inline.find(qn("wp:docPr"))
    frame_pr = inline.find(qn("wp:cNvGraphicFramePr"))
    graphic = inline.find(qn("a:graphic"))

    emu = int(dist_pt * 12700)
    anchor = OxmlElement("wp:anchor")
    for k, v in {"distT": "0", "distB": "0", "distL": str(emu), "distR": str(emu),
                 "simplePos": "0", "relativeHeight": "2", "behindDoc": "0",
                 "locked": "0", "layoutInCell": "1", "allowOverlap": "1"}.items():
        anchor.set(k, v)

    sp = OxmlElement("wp:simplePos")
    sp.set("x", "0")
    sp.set("y", "0")
    anchor.append(sp)

    ph = OxmlElement("wp:positionH")
    ph.set("relativeFrom", "margin")
    al = OxmlElement("wp:align")
    al.text = side
    ph.append(al)
    anchor.append(ph)

    pv = OxmlElement("wp:positionV")
    pv.set("relativeFrom", "paragraph")
    off = OxmlElement("wp:posOffset")
    off.text = "0"
    pv.append(off)
    anchor.append(pv)

    anchor.append(extent)

    ee = OxmlElement("wp:effectExtent")
    for k in "ltrb":
        ee.set(k, "0")
    anchor.append(ee)

    wrap = OxmlElement("wp:wrapSquare")          # <- thứ làm chữ trôi quanh ảnh
    wrap.set("wrapText", "bothSides")
    anchor.append(wrap)

    anchor.append(doc_pr)
    if frame_pr is not None:
        anchor.append(frame_pr)
    anchor.append(graphic)

    drawing.remove(inline)
    drawing.append(anchor)


# ------------------------------------------------------------ dựng nội dung ---

def add_inline(par, nodes, figures):
    for n in nodes:
        t = n["type"]
        if t == "text":
            r = par.add_run(n["text"])
            marks = n.get("marks", [])
            r.bold = "bold" in marks
            r.italic = "italic" in marks
            r.underline = "underline" in marks
        elif t == "math":
            # TODO: dựng MathML cho Word. Hiện in TeX dạng chữ nghiêng.
            r = par.add_run(n["tex"])
            r.italic = True
            r.font.name = "Cambria Math"
        elif t == "hard_break":
            par.add_run().add_break()
        elif t == "image_inline":
            row = (figures or {}).get(n["figure_id"])
            if row and row.get("storage_path"):
                par.add_run().add_picture(
                    row["storage_path"],
                    width=Inches(TEXT_WIDTH_IN * float(row.get("width") or 0.45)))


def add_blocks(host, nodes, figures):
    for n in nodes:
        t = n["type"]
        if t == "paragraph":
            add_inline(host.add_paragraph(), n["content"], figures)
        elif t == "math_block":
            p = host.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run(n["tex"]).italic = True
        elif t == "image":
            p = host.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_inline(p, [{"type": "image_inline", "figure_id": n["figure_id"]}], figures)
        elif t == "table":
            add_table(host, n, figures)
        elif t == "list":
            style = "List Number" if n.get("ordered") else "List Bullet"
            for item in n["items"]:
                for b in item:
                    if b["type"] == "paragraph":
                        add_inline(host.add_paragraph(style=style), b["content"], figures)


def add_table(host, node, figures):
    ncol = max(len(r) for r in node["rows"])
    t = host.add_table(rows=0, cols=ncol)
    t.style = "Table Grid"
    for row in node["rows"]:
        cells = t.add_row().cells
        for i, c in enumerate(row):
            if i < ncol:
                cells[i].text = ""
                add_inline(cells[i].paragraphs[0], c["content"], figures)


def write_field(doc, tree, figures, lead=None):
    """Ghi ruột một trường. `lead` ("Câu 3: ") ghép vào ĐÚNG đoạn đầu của đề,
    để tiêu đề và đề bài nằm chung một dòng — và vì đoạn đầu cũng là đoạn mang
    mỏ neo của ảnh, tách ra thì chữ không trôi quanh nữa."""
    side = tree.get("side", "center")
    floating = side in ("left", "right")
    blocks = [n for n in tree.get("content", [])
              if not (floating and n["type"] == "image")]
    figs = [n for n in tree.get("content", []) if n["type"] == "image"] if floating else []

    first = None
    if (blocks and blocks[0]["type"] == "paragraph") or lead or figs:
        first = doc.add_paragraph()

    if first is not None:
        for n in figs:
            row = (figures or {}).get(n["figure_id"])
            if row and row.get("storage_path"):
                run = first.add_run()
                run.add_picture(row["storage_path"],
                                width=Inches(TEXT_WIDTH_IN * float(row.get("width") or 0.45)))
                float_picture(run, side=side)
        if lead:
            first.add_run(lead).bold = True
        if blocks and blocks[0]["type"] == "paragraph":
            add_inline(first, blocks[0]["content"], figures)
            blocks = blocks[1:]

    add_blocks(doc, blocks, figures)


def add_question(doc, rec, num, figures=None):
    """Một câu -> các đoạn trong tài liệu Word."""
    qt = rec["question_type"]
    mark = len(doc.paragraphs)

    write_field(doc, rec["content_doc"], figures, lead=f"Câu {num}: ")

    for i, o in enumerate(rec.get("options") or []):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.25)
        label = ("Đúng " if o["is_correct"] else "Sai ") if qt == "tf" \
            else f"{'ABCD'[i % 4]}. "
        p.add_run(label).bold = True
        for b in o["content_doc"]["content"]:
            if b["type"] == "paragraph":
                add_inline(p, b["content"], figures)

    sol = rec.get("solution_doc") or {}
    if sol.get("content"):
        write_field(doc, sol, figures, lead="Lời giải. ")

    # Giữ cả cụm câu trên cùng một trang. Không có cái này thì ảnh neo vẫn đúng
    # nhưng Word được phép cắt ngang câu, khiến mấy phương án cuối rơi sang
    # trang sau và rời khỏi hình.
    if rec["content_doc"].get("side", "center") != "center":
        for p in doc.paragraphs[mark:-1]:
            p.paragraph_format.keep_with_next = True

    doc.add_paragraph()


def to_docx(recs, path, figures=None, title=None):
    """Danh sách câu -> file `.docx`."""
    doc = Document()
    if title:
        doc.add_heading(title, level=1)
    for i, rec in enumerate(recs, 1):
        add_question(doc, rec, i, figures)
    doc.save(path)
    return path
